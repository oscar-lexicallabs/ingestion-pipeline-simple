from docx import Document
import pypandoc
import os
import random
import string
import argparse

WORDS = "./words_alpha.txt"

def make_new_doc(dir: str) -> None:
    doc = Document()
    with open(WORDS, mode="rt") as f:
        contents = f.read().splitlines()

        file_name = "_".join([
            random.choice(contents)
            for _ in range(random.randint(2,5))
        ])
        doc.add_heading(file_name.replace("_"," ").title())

        num_paragraphs = random.randint(5,10)
        for _ in range(num_paragraphs):
            para_len = random.randint(30,100)
            para = " ".join([
                random.choice(contents)
                for _ in range(para_len)
            ])
            doc.add_paragraph(para)
    
    rand_bits = "".join([
        random.choice(string.ascii_lowercase)
        for _ in range(4)
    ])
    save_name = file_name + "_" + rand_bits + ".docx"
    save_path = os.path.join(dir, save_name)
    doc.save(save_path)

def convert_to_pdf(dir: str) -> None:
    assert os.path.isdir(dir)
    for file in os.listdir(dir):
        if file.endswith(".docx"):
            pdf_file = file.replace(".docx",".pdf")
            pypandoc.convert_file(
                os.path.join(dir, file),
                "pdf",
                outputfile=os.path.join(dir, pdf_file)
            )

def clean_dir(dir: str) -> None:
    assert os.path.isdir(dir)
    for file in os.listdir(dir):
        if file.endswith(".docx") or file.endswith(".pdf"):
            os.remove(os.path.join(dir, file))

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("-m", "--make", action="store_true")
    parser.add_argument("-n", "--num_to_make", type=int, default=1)
    parser.add_argument("-c", "--clean", action="store_true")
    parser.add_argument("-d", "--dir", type=str, default=".")
    opt = parser.parse_args()
    if opt.make:
        for _ in range(opt.num_to_make):
            make_new_doc(opt.dir)
        convert_to_pdf(opt.dir)
    if opt.clean:
        clean_dir(opt.dir)